"""
Text extraction service for PrepAble.

Extracts raw text from uploaded resume files (PDF, DOCX, TXT).
All returned text is a single clean string — callers do not need to
know which parser was used.
"""

import io
from pathlib import Path
from typing import Union


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract readable text from a resume file.

    Supported formats:
        .pdf  — via PyPDF2 (handles multi-page; skips pages that yield no text)
        .docx — via python-docx (paragraphs + table cells; covers skills/experience tables)
        .txt  — decoded as UTF-8, falling back to latin-1

    Raises:
        ValueError  — unsupported extension or empty file
        RuntimeError — file is corrupt / unreadable
    """
    if not file_bytes:
        raise ValueError("File is empty")

    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    elif suffix == ".docx":
        return _extract_docx(file_bytes)
    elif suffix == ".txt":
        return _extract_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{suffix}'. Upload PDF, DOCX, or TXT.")


# ---------------------------------------------------------------------------
# Internal parsers
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2  # noqa: PLC0415 — optional heavy import
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is not installed") from exc

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise RuntimeError(f"Could not open PDF: {exc}") from exc

    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        except Exception:  # noqa: BLE001 — skip unreadable pages, don't crash
            continue

    if not pages:
        raise ValueError(
            "PDF appears to be empty or contains no extractable text "
            "(scanned image PDFs are not supported)"
        )

    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise RuntimeError(f"Could not open DOCX: {exc}") from exc

    parts = []

    # Extract paragraph text (body, headings, bullet points)
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Extract table cell text — critical for resumes that use tables to
    # lay out skills, technologies, or experience columns. python-docx's
    # doc.paragraphs does NOT include table cells, so this pass is required.
    for table in doc.tables:
        for row in table.rows:
            # Deduplicate merged cells (python-docx repeats merged cell text)
            seen: set[str] = set()
            row_cells = []
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    seen.add(text)
                    row_cells.append(text)
            if row_cells:
                parts.append(" | ".join(row_cells))

    if not parts:
        raise ValueError("DOCX file appears to be empty")

    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                return text
        except (UnicodeDecodeError, ValueError):
            continue

    raise ValueError("TXT file could not be decoded (try saving as UTF-8)")
